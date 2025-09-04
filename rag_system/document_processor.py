#!/usr/bin/env python3
"""
Advanced RAG System with Document Upload and Processing
Supports multiple document formats, chunking strategies, and vector stores
"""

import os
import torch
import numpy as np
from typing import List, Dict, Optional, Union, Tuple
from pathlib import Path
import json
import pickle
from datetime import datetime
import hashlib
import logging
from dataclasses import dataclass, asdict
from abc import ABC, abstractmethod

# Document processing
import PyPDF2
import docx
from bs4 import BeautifulSoup
import pandas as pd
from PIL import Image
import pytesseract
import markdown

# Vector operations
import faiss
from sentence_transformers import SentenceTransformer
import chromadb
from sklearn.metrics.pairwise import cosine_similarity

# Text processing
from langchain.text_splitter import (
    RecursiveCharacterTextSplitter,
    CharacterTextSplitter,
    TokenTextSplitter
)
from langchain.docstore.document import Document as LangchainDocument

# Web search
import requests
from bs4 import BeautifulSoup

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


@dataclass
class RAGConfig:
    """Configuration for the RAG system"""

    # Embedding model settings
    embedding_model: str = "all-MiniLM-L6-v2"
    device: str = "cuda" if torch.cuda.is_available() else "cpu"

    # Vector store settings
    vector_store_type: str = "faiss"  # Options: faiss, chromadb, simple
    vector_store_path: str = "./vector_store"

    # Chunking settings
    chunk_size: int = 1000
    chunk_overlap: int = 200
    chunking_strategy: str = "recursive"  # Options: recursive, character, token

    # Retrieval settings
    top_k: int = 5
    similarity_threshold: float = 0.3

    # Processing settings
    max_file_size_mb: int = 50
    supported_formats: List[str] = None

    def __post_init__(self):
        if self.supported_formats is None:
            self.supported_formats = ['.pdf', '.docx', '.txt', '.md', '.html', '.csv', '.xlsx', '.jpg', '.png']


class DocumentProcessor:
    """Process various document formats"""

    @staticmethod
    def extract_text_from_pdf(file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {e}")
            return ""

    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {e}")
            return ""

    @staticmethod
    def extract_text_from_html(file_path: str) -> str:
        """Extract text from HTML file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                soup = BeautifulSoup(file.read(), 'html.parser')
                return soup.get_text()
        except Exception as e:
            logger.error(f"Error processing HTML {file_path}: {e}")
            return ""

    @staticmethod
    def extract_text_from_csv(file_path: str) -> str:
        """Extract text from CSV file"""
        try:
            df = pd.read_csv(file_path)
            return df.to_string()
        except Exception as e:
            logger.error(f"Error processing CSV {file_path}: {e}")
            return ""

    @staticmethod
    def extract_text_from_excel(file_path: str) -> str:
        """Extract text from Excel file"""
        try:
            df = pd.read_excel(file_path)
            return df.to_string()
        except Exception as e:
            logger.error(f"Error processing Excel {file_path}: {e}")
            return ""

    @staticmethod
    def extract_text_from_image(file_path: str) -> str:
        """Extract text from image using OCR"""
        try:
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image)
            return text
        except Exception as e:
            logger.error(f"Error processing image {file_path}: {e}")
            return ""

    @staticmethod
    def extract_text_from_markdown(file_path: str) -> str:
        """Extract text from Markdown file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                md_content = file.read()
                html = markdown.markdown(md_content)
                soup = BeautifulSoup(html, 'html.parser')
                return soup.get_text()
        except Exception as e:
            logger.error(f"Error processing Markdown {file_path}: {e}")
            return ""

    @classmethod
    def process_file(cls, file_path: str) -> Tuple[str, Dict]:
        """Process a file and extract text with metadata"""
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        # Get file metadata
        metadata = {
            "filename": file_path.name,
            "file_path": str(file_path),
            "file_size": file_path.stat().st_size,
            "file_extension": file_path.suffix.lower(),
            "processed_at": datetime.now().isoformat(),
            "file_hash": cls._get_file_hash(file_path)
        }

        # Extract text based on file type
        extension = file_path.suffix.lower()

        if extension == '.pdf':
            text = cls.extract_text_from_pdf(str(file_path))
        elif extension == '.docx':
            text = cls.extract_text_from_docx(str(file_path))
        elif extension == '.html':
            text = cls.extract_text_from_html(str(file_path))
        elif extension == '.csv':
            text = cls.extract_text_from_csv(str(file_path))
        elif extension in ['.xlsx', '.xls']:
            text = cls.extract_text_from_excel(str(file_path))
        elif extension in ['.jpg', '.jpeg', '.png', '.tiff', '.bmp']:
            text = cls.extract_text_from_image(str(file_path))
        elif extension == '.md':
            text = cls.extract_text_from_markdown(str(file_path))
        elif extension == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
        else:
            raise ValueError(f"Unsupported file format: {extension}")

        metadata["text_length"] = len(text)
        metadata["word_count"] = len(text.split())

        return text, metadata

    @staticmethod
    def _get_file_hash(file_path: Path) -> str:
        """Generate hash for file content"""
        hasher = hashlib.md5()
        with open(file_path, 'rb') as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hasher.update(chunk)
        return hasher.hexdigest()


class TextChunker:
    """Advanced text chunking with multiple strategies"""

    def __init__(self, config: RAGConfig):
        self.config = config
        self._setup_splitter()

    def _setup_splitter(self):
        """Setup text splitter based on strategy"""
        if self.config.chunking_strategy == "recursive":
            self.splitter = RecursiveCharacterTextSplitter(
                chunk_size=self.config.chunk_size,
                chunk_overlap=self.config.chunk_overlap,
                length_function=len,
                separators=["\n\n", "\n", " ", ""]
            )
        elif self.config.chunking_strategy == "character":
            self.splitter = CharacterTextSplitter(
                chunk_size=self.config.chunk_size,
                chunk_overlap=self.config.chunk_overlap,
                length_function=len
            )
        elif self.config.chunking_strategy == "token":
            self.splitter = TokenTextSplitter(
                chunk_size=self.config.chunk_size,
                chunk_overlap=self.config.chunk_overlap
            )
        else:
            raise ValueError(f"Unknown chunking strategy: {self.config.chunking_strategy}")

    def chunk_text(self, text: str, metadata: Dict) -> List[Dict]:
        """Chunk text and return documents with metadata"""

        # Create langchain document
        doc = LangchainDocument(page_content=text, metadata=metadata)

        # Split document
        chunks = self.splitter.split_documents([doc])

        # Convert to our format
        chunked_docs = []
        for i, chunk in enumerate(chunks):
            chunk_metadata = chunk.metadata.copy()
            chunk_metadata.update({
                "chunk_id": i,
                "chunk_length": len(chunk.page_content),
                "total_chunks": len(chunks)
            })

            chunked_docs.append({
                "content": chunk.page_content,
                "metadata": chunk_metadata
            })

        return chunked_docs


class VectorStore(ABC):
    """Abstract base class for vector stores"""

    @abstractmethod
    def add_documents(self, documents: List[Dict], embeddings: np.ndarray):
        pass

    @abstractmethod
    def search(self, query_embedding: np.ndarray, k: int) -> List[Tuple[int, float]]:
        pass

    @abstractmethod
    def save(self, path: str):
        pass

    @abstractmethod
    def load(self, path: str):
        pass


class FAISSVectorStore(VectorStore):
    """FAISS-based vector store"""

    def __init__(self, dimension: int, use_gpu: bool = False):
        self.dimension = dimension
        self.use_gpu = use_gpu and torch.cuda.is_available()

        if self.use_gpu:
            res = faiss.StandardGpuResources()
            self.index = faiss.GpuIndexFlatIP(res, dimension)
        else:
            self.index = faiss.IndexFlatIP(dimension)

        self.documents = []

    def add_documents(self, documents: List[Dict], embeddings: np.ndarray):
        """Add documents and their embeddings"""
        self.documents.extend(documents)
        self.index.add(embeddings.astype(np.float32))

    def search(self, query_embedding: np.ndarray, k: int) -> List[Tuple[int, float]]:
        """Search for similar documents"""
        scores, indices = self.index.search(
            query_embedding.astype(np.float32),
            min(k, self.index.ntotal)
        )

        results = []
        for idx, score in zip(indices[0], scores[0]):
            if idx != -1:  # FAISS returns -1 for empty slots
                results.append((int(idx), float(score)))

        return results

    def save(self, path: str):
        """Save index and documents"""
        os.makedirs(path, exist_ok=True)

        # Save FAISS index
        faiss.write_index(self.index, os.path.join(path, "faiss.index"))

        # Save documents
        with open(os.path.join(path, "documents.pkl"), 'wb') as f:
            pickle.dump(self.documents, f)

    def load(self, path: str):
        """Load index and documents"""
        # Load FAISS index
        self.index = faiss.read_index(os.path.join(path, "faiss.index"))

        # Load documents
        with open(os.path.join(path, "documents.pkl"), 'rb') as f:
            self.documents = pickle.load(f)


class ChromaVectorStore(VectorStore):
    """ChromaDB-based vector store"""

    def __init__(self, collection_name: str = "rag_documents", persist_directory: str = "./chroma_db"):
        self.client = chromadb.PersistentClient(path=persist_directory)
        self.collection = self.client.get_or_create_collection(name=collection_name)
        self.documents = []

    def add_documents(self, documents: List[Dict], embeddings: np.ndarray):
        """Add documents and their embeddings"""
        ids = [f"doc_{len(self.documents) + i}" for i in range(len(documents))]

        self.collection.add(
            embeddings=embeddings.tolist(),
            documents=[doc["content"] for doc in documents],
            metadatas=[doc["metadata"] for doc in documents],
            ids=ids
        )

        self.documents.extend(documents)

    def search(self, query_embedding: np.ndarray, k: int) -> List[Tuple[int, float]]:
        """Search for similar documents"""
        results = self.collection.query(
            query_embeddings=[query_embedding.tolist()],
            n_results=min(k, len(self.documents))
        )

        # Convert to our format (ChromaDB doesn't return indices directly)
        search_results = []
        for i, (doc_id, distance) in enumerate(zip(results['ids'][0], results['distances'][0])):
            # ChromaDB returns distances, convert to similarity scores
            similarity = 1 - distance
            search_results.append((i, float(similarity)))

        return search_results

    def save(self, path: str):
        """ChromaDB auto-persists"""
        pass

    def load(self, path: str):
        """ChromaDB auto-loads"""
        pass


class AdvancedRAGSystem:
    """Advanced RAG system with document upload and processing"""

    def __init__(self, config: RAGConfig):
        self.config = config
        self.device = config.device

        # Initialize components
        self.embedder = SentenceTransformer(config.embedding_model, device=self.device)
        self.text_chunker = TextChunker(config)

        # Initialize vector store
        self.vector_store = self._setup_vector_store()

        # Document tracking
        self.processed_files = set()
        self.document_metadata = []

        logger.info(f"Initialized AdvancedRAGSystem with {config.vector_store_type} vector store")

    def _setup_vector_store(self) -> VectorStore:
        """Setup vector store based on configuration"""
        embedding_dim = self.embedder.get_sentence_embedding_dimension()

        if self.config.vector_store_type == "faiss":
            return FAISSVectorStore(
                dimension=embedding_dim,
                use_gpu=(self.device == "cuda")
            )
        elif self.config.vector_store_type == "chromadb":
            return ChromaVectorStore(persist_directory=self.config.vector_store_path)
        else:
            raise ValueError(f"Unknown vector store type: {self.config.vector_store_type}")

    def process_uploaded_files(self, file_paths: List[str]) -> Dict[str, any]:
        """Process multiple uploaded files"""
        results = {
            "processed_files": [],
            "failed_files": [],
            "total_chunks": 0,
            "processing_time": 0
        }

        start_time = datetime.now()

        all_documents = []
        all_embeddings = []

        for file_path in file_paths:
            try:
                # Check if file was already processed
                file_hash = DocumentProcessor._get_file_hash(Path(file_path))
                if file_hash in self.processed_files:
                    logger.info(f"File already processed: {file_path}")
                    continue

                # Check file size
                file_size_mb = Path(file_path).stat().st_size / (1024 * 1024)
                if file_size_mb > self.config.max_file_size_mb:
                    results["failed_files"].append({
                        "file": file_path,
                        "reason": f"File too large: {file_size_mb:.1f}MB > {self.config.max_file_size_mb}MB"
                    })
                    continue

                # Process file
                text, metadata = DocumentProcessor.process_file(file_path)

                if not text.strip():
                    results["failed_files"].append({
                        "file": file_path,
                        "reason": "No text content extracted"
                    })
                    continue

                # Chunk text
                chunked_docs = self.text_chunker.chunk_text(text, metadata)

                # Generate embeddings
                contents = [doc["content"] for doc in chunked_docs]
                embeddings = self.embedder.encode(
                    contents,
                    batch_size=32,
                    show_progress_bar=True,
                    convert_to_numpy=True,
                    normalize_embeddings=True
                )

                # Add to collections
                all_documents.extend(chunked_docs)
                all_embeddings.append(embeddings)

                # Track processed file
                self.processed_files.add(file_hash)
                self.document_metadata.append(metadata)

                results["processed_files"].append({
                    "file": file_path,
                    "chunks": len(chunked_docs),
                    "text_length": len(text)
                })
                results["total_chunks"] += len(chunked_docs)

                logger.info(f"Processed {file_path}: {len(chunked_docs)} chunks")

            except Exception as e:
                logger.error(f"Error processing {file_path}: {e}")
                results["failed_files"].append({
                    "file": file_path,
                    "reason": str(e)
                })

        # Add all documents to vector store
        if all_documents:
            combined_embeddings = np.vstack(all_embeddings)
            self.vector_store.add_documents(all_documents, combined_embeddings)

            # Save vector store
            self.save_vector_store()

        results["processing_time"] = (datetime.now() - start_time).total_seconds()
        return results

    def add_text_documents(self, texts: List[str], metadatas: List[Dict] = None) -> int:
        """Add raw text documents to the system"""
        if metadatas is None:
            metadatas = [{"source": "text_input", "index": i} for i in range(len(texts))]

        all_documents = []
        all_embeddings = []

        for text, metadata in zip(texts, metadatas):
            # Chunk text
            chunked_docs = self.text_chunker.chunk_text(text, metadata)

            # Generate embeddings
            contents = [doc["content"] for doc in chunked_docs]
            embeddings = self.embedder.encode(
                contents,
                batch_size=32,
                convert_to_numpy=True,
                normalize_embeddings=True
            )

            all_documents.extend(chunked_docs)
            all_embeddings.append(embeddings)

        # Add to vector store
        if all_documents:
            combined_embeddings = np.vstack(all_embeddings)
            self.vector_store.add_documents(all_documents, combined_embeddings)
            self.save_vector_store()

        return len(all_documents)

    def retrieve_documents(self, query: str, k: int = None) -> List[Dict]:
        """Retrieve relevant documents for a query"""
        if k is None:
            k = self.config.top_k

        # Generate query embedding
        query_embedding = self.embedder.encode(
            [query],
            convert_to_numpy=True,
            normalize_embeddings=True
        )

        # Search vector store
        search_results = self.vector_store.search(query_embedding, k)

        # Filter by similarity threshold and return documents
        relevant_docs = []
        for idx, score in search_results:
            if score >= self.config.similarity_threshold:
                doc = self.vector_store.documents[idx].copy()
                doc["similarity_score"] = score
                relevant_docs.append(doc)

        return relevant_docs

    def web_search_enhanced(self, query: str, num_results: int = 5) -> List[Dict]:
        """Enhanced web search with better parsing"""
        try:
            # Use multiple search engines
            results = []

            # DuckDuckGo API
            duckduckgo_results = self._search_duckduckgo(query, num_results)
            results.extend(duckduckgo_results)

            # You can add more search engines here

            return results[:num_results]

        except Exception as e:
            logger.error(f"Web search error: {e}")
            return []

    def _search_duckduckgo(self, query: str, num_results: int) -> List[Dict]:
        """Search using DuckDuckGo API"""
        try:
            url = f"https://api.duckduckgo.com/?q={query}&format=json&no_html=1&skip_disambig=1"
            response = requests.get(url, timeout=10)
            data = response.json()

            results = []

            # Abstract
            if data.get('Abstract'):
                results.append({
                    'title': data.get('AbstractSource', 'DuckDuckGo'),
                    'content': data['Abstract'],
                    'url': data.get('AbstractURL', ''),
                    'source': 'duckduckgo',
                    'relevance': 1.0
                })

            # Related topics
            for topic in data.get('RelatedTopics', [])[:num_results - 1]:
                if isinstance(topic, dict) and topic.get('Text'):
                    results.append({
                        'title': topic.get('FirstURL', '').split('/')[-1] or 'Related',
                        'content': topic['Text'],
                        'url': topic.get('FirstURL', ''),
                        'source': 'duckduckgo',
                        'relevance': 0.8
                    })

            return results

        except Exception as e:
            logger.error(f"DuckDuckGo search error: {e}")
            return []

    def generate_response(self, query: str, retrieved_docs: List[Dict],
                          search_results: List[Dict] = None) -> str:
        """Generate comprehensive response from retrieved information"""

        context_parts = []

        # Add retrieved documents
        if retrieved_docs:
            context_parts.append("📚 **From Knowledge Base:**")
            for i, doc in enumerate(retrieved_docs[:3], 1):
                filename = doc['metadata'].get('filename', 'Unknown')
                similarity = doc.get('similarity_score', 0)
                context_parts.append(
                    f"{i}. **{filename}** (similarity: {similarity:.2f})\n"
                    f"   {doc['content'][:300]}{'...' if len(doc['content']) > 300 else ''}"
                )

        # Add web search results
        if search_results:
            context_parts.append("\n🌐 **From Web Search:**")
            for i, result in enumerate(search_results[:3], 1):
                context_parts.append(
                    f"{i}. **{result['title']}**\n"
                    f"   {result['content'][:200]}{'...' if len(result['content']) > 200 else ''}\n"
                    f"   Source: {result['url']}"
                )

        if not context_parts:
            return "❌ No relevant information found. Try rephrasing your question or adding more documents to the knowledge base."

        # Generate structured response
        response = f"""🎯 **Query:** {query}

{chr(10).join(context_parts)}

💡 **Analysis:** Based on the retrieved information from both the knowledge base and web sources, this provides a comprehensive answer to your question. The similarity scores indicate how relevant each document is to your query."""

        return response

    def process_query_complete(self, query: str, include_web_search: bool = True) -> Dict:
        """Complete query processing pipeline"""
        start_time = datetime.now()

        # Retrieve documents
        retrieved_docs = self.retrieve_documents(query)

        # Web search if enabled
        search_results = []
        if include_web_search:
            search_results = self.web_search_enhanced(query)

        # Generate response
        response = self.generate_response(query, retrieved_docs, search_results)

        processing_time = (datetime.now() - start_time).total_seconds()

        return {
            "query": query,
            "response": response,
            "retrieved_documents": len(retrieved_docs),
            "web_results": len(search_results),
            "processing_time": processing_time,
            "retrieved_docs": retrieved_docs,
            "search_results": search_results
        }

    def get_system_stats(self) -> Dict:
        """Get system statistics"""
        return {
            "total_documents": len(self.vector_store.documents),
            "processed_files": len(self.processed_files),
            "embedding_model": self.config.embedding_model,
            "vector_store_type": self.config.vector_store_type,
            "device": self.device,
            "chunk_size": self.config.chunk_size,
            "chunk_overlap": self.config.chunk_overlap
        }

    def save_vector_store(self):
        """Save the vector store"""
        self.vector_store.save(self.config.vector_store_path)

        # Save metadata
        metadata_path = os.path.join(self.config.vector_store_path, "metadata.json")
        with open(metadata_path, 'w') as f:
            json.dump({
                "config": asdict(self.config),
                "processed_files": list(self.processed_files),
                "document_metadata": self.document_metadata,
                "last_updated": datetime.now().isoformat()
            }, f, indent=2)

    def load_vector_store(self):
        """Load the vector store"""
        if os.path.exists(self.config.vector_store_path):
            self.vector_store.load(self.config.vector_store_path)

            # Load metadata
            metadata_path = os.path.join(self.config.vector_store_path, "metadata.json")
            if os.path.exists(metadata_path):
                with open(metadata_path, 'r') as f:
                    metadata = json.load(f)
                    self.processed_files = set(metadata.get("processed_files", []))
                    self.document_metadata = metadata.get("document_metadata", [])

            logger.info(f"Loaded vector store with {len(self.vector_store.documents)} documents")
        else:
            logger.info("No existing vector store found")

    def clear_documents(self):
        """Clear all documents from the system"""
        self.vector_store = self._setup_vector_store()
        self.processed_files.clear()
        self.document_metadata.clear()

        # Remove saved files
        if os.path.exists(self.config.vector_store_path):
            import shutil
            shutil.rmtree(self.config.vector_store_path)

        logger.info("Cleared all documents from the system")


def create_rag_system(config: RAGConfig = None) -> AdvancedRAGSystem:
    """Factory function to create RAG system"""
    if config is None:
        config = RAGConfig()

    system = AdvancedRAGSystem(config)
    system.load_vector_store()  # Load existing data if available

    return system


# Example usage
if __name__ == "__main__":
    # Create RAG system
    config = RAGConfig(
        embedding_model="all-MiniLM-L6-v2",
        vector_store_type="faiss",
        chunk_size=1000,
        chunk_overlap=200,
        top_k=5
    )

    rag_system = create_rag_system(config)

    # Example: Process files
    # files = ["document1.pdf", "document2.docx", "data.csv"]
    # results = rag_system.process_uploaded_files(files)
    # print(f"Processed {results['total_chunks']} chunks from {len(results['processed_files'])} files")

    # Example: Add text documents
    texts = [
        "Machine learning is a subset of artificial intelligence that focuses on algorithms.",
        "Deep learning uses neural networks with multiple layers to process data."
    ]
    rag_system.add_text_documents(texts)

    # Example: Query
    query = "What is machine learning?"
    result = rag_system.process_query_complete(query)
    print(f"Query: {result['query']}")
    print(f"Response: {result['response']}")
    print(f"Processing time: {result['processing_time']:.2f}s")